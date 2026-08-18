#!/usr/bin/env python3
"""Build the three Week 6 Word reports from machine-readable evidence only."""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
import json
from pathlib import Path
import re
from typing import Any, Iterable

REPORT_FILENAMES = (
    "完整测试与覆盖率报告.docx",
    "性能优化基准报告.docx",
    "缺陷修复与本地数据安全审查报告.docx",
)
COMMIT = re.compile(r"^[0-9a-f]{40}$")


def validate_report_inputs(manifest: dict[str, Any], *, draft: bool) -> dict[str, Any]:
    source_commit = manifest.get("source_commit")
    if not isinstance(source_commit, str) or not COMMIT.fullmatch(source_commit):
        raise ValueError("manifest source_commit must be a full Git commit")
    gates = manifest.get("gates")
    if not isinstance(gates, list):
        raise ValueError("manifest gates must be a list")
    statuses = {gate.get("gate_id"): gate.get("status") for gate in gates if isinstance(gate, dict)}
    missing = [f"G{i}" for i in range(10) if statuses.get(f"G{i}") != "PASS"]
    if missing and not draft:
        raise ValueError(f"final reports require PASS gates; incomplete: {', '.join(missing)}")
    return {
        "source_commit": source_commit,
        "document_status": "DRAFT - NOT ACCEPTED" if missing else "FINAL - ACCEPTED",
        "incomplete_gates": missing,
        "gate_statuses": statuses,
    }


def _read_optional(path: Path | None) -> dict[str, Any]:
    if path is None or not path.is_file():
        return {}
    value = json.loads(path.read_text(encoding="utf-8"))
    return value if isinstance(value, dict) else {}


def _set_cell_margins(cell: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _apply_table_geometry(table: Any, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _font_run(run: Any, *, size: float | None = None, bold: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _style_document(document: Any, *, title: str, subtitle: str, status: str, commit: str) -> None:
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = document.styles
    for name, size, before, after, bold in (
        ("Normal", 11, 0, 6, False),
        ("Title", 23, 0, 4, True),
        ("Subtitle", 13, 0, 14, False),
        ("Heading 1", 16, 16, 8, True),
        ("Heading 2", 13, 12, 6, True),
        ("Heading 3", 12, 8, 4, True),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        borders = style._element.get_or_add_pPr().find(qn("w:pBdr"))
        if borders is not None:
            style._element.get_or_add_pPr().remove(borders)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _font_run(header.add_run("第六周系统集成验收"), size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font_run(footer.add_run(f"Source {commit[:12]}"), size=9)
    title_p = document.add_paragraph(style="Title")
    _font_run(title_p.add_run(title), size=23, bold=True)
    subtitle_p = document.add_paragraph(style="Subtitle")
    _font_run(subtitle_p.add_run(subtitle), size=13)
    meta = document.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    values = (("文档状态", status), ("候选提交", commit), ("生成时间", datetime.now(timezone.utc).isoformat()))
    for row, pair in zip(meta.rows, values, strict=True):
        for cell, text in zip(row.cells, pair, strict=True):
            cell.text = text
    _apply_table_geometry(meta, [1800, 7560])


def _add_table(document: Any, headers: Iterable[str], rows: Iterable[Iterable[Any]], widths: list[int]) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    values = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(widths))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        cell.text = str(text)
        for run in cell.paragraphs[0].runs:
            _font_run(run, bold=True)
    for values_row in values:
        cells = table.add_row().cells
        for cell, text in zip(cells, values_row, strict=True):
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _apply_table_geometry(table, widths)


def _add_gate_summary(document: Any, context: dict[str, Any]) -> None:
    document.add_heading("验收门禁总览", level=1)
    _add_table(
        document,
        ("门禁", "状态"),
        ((f"G{i}", context["gate_statuses"].get(f"G{i}", "MISSING")) for i in range(10)),
        [1800, 7560],
    )
    if context["incomplete_gates"]:
        paragraph = document.add_paragraph()
        _font_run(
            paragraph.add_run("本报告为未完成草稿，不得作为验收通过证明。未完成门禁：" + ", ".join(context["incomplete_gates"])),
            bold=True,
        )


def _build_test_report(output: Path, context: dict[str, Any], coverage: dict[str, Any], stress: dict[str, Any]) -> None:
    from docx import Document

    doc = Document()
    _style_document(doc, title="完整测试与覆盖率报告", subtitle="单元、集成、端到端与压力测试", status=context["document_status"], commit=context["source_commit"])
    _add_gate_summary(doc, context)
    doc.add_page_break()
    doc.add_heading("测试分层", level=1)
    _add_table(doc, ("层级", "验收范围", "结果"), (
        ("单元", "核心 Python 模块；仅 unit 标记", coverage.get("status", "EVIDENCE PENDING")),
        ("集成", "API、Tika、Chroma、真实依赖", context["gate_statuses"].get("G4")),
        ("端到端", "Flutter UI 至真实后端全链路", context["gate_statuses"].get("G2")),
        ("压力", "10k 记录、500 查询、30 分钟 soak", stress.get("status", "EVIDENCE PENDING")),
    ), [1400, 5560, 2400])
    doc.add_heading("覆盖率判定", level=1)
    percent = coverage.get("totals", {}).get("percent_covered") if isinstance(coverage.get("totals"), dict) else None
    doc.add_paragraph(f"核心模块合并语句覆盖率（未四舍五入）：{percent if percent is not None else '证据待生成'}；门槛为 90.00%。")
    doc.add_heading("结论与限制", level=1)
    doc.add_paragraph("最终结论只由 G0-G9 严格证据清单决定。任何必需测试跳过、失败、BLOCKED 或证据缺失均不得写为通过。")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def _performance_summary_rows(comparison: dict[str, Any]) -> list[tuple[str, str, str]]:
    candidate = comparison.get("candidate_medians", {})
    improvements = comparison.get("improvements_percent", {})
    if not isinstance(candidate, dict) or not isinstance(improvements, dict):
        return [("性能比较", "证据待生成", "证据待生成")]
    definitions = (
        ("混合负载嵌入 P95", "embedding_combined_p95_ms", "ms"),
        ("混合负载向量查询 P95", "vector_query_p95_ms", "ms"),
        ("热缓存嵌入 P95", "embedding_hot_p95_ms", "ms"),
        ("热缓存向量查询 P95", "vector_query_hot_p95_ms", "ms"),
        ("峰值 RSS", "peak_rss_bytes", "bytes"),
        ("全链路检索 P95", "full_search_p95_ms", "ms"),
    )
    rows: list[tuple[str, str, str]] = []
    for label, key, unit in definitions:
        value = candidate.get(key)
        if isinstance(value, (int, float)):
            measured = f"{value / 1024 / 1024:.2f} MiB" if unit == "bytes" else f"{value:.3f} ms"
        else:
            measured = "证据待生成"
        improvement = improvements.get(key)
        change = f"{improvement:.3f}%" if isinstance(improvement, (int, float)) else "证据待生成"
        rows.append((label, measured, change))
    return rows


def _security_summary_rows(security: dict[str, Any]) -> list[tuple[str, str, str]]:
    checks = security.get("checks")
    isolation = security.get("network_isolation")
    if not isinstance(checks, dict) or not isinstance(isolation, dict):
        return [("安全审查", "证据待生成", "证据待生成")]
    isolation_status = checks.get("network_isolation", "PASS" if isolation.get("enforced") is True else "FAIL")
    rows = [
        (
            "网络隔离",
            f"{isolation.get('method', '未知')}；{isolation.get('sample_seconds', '未知')} 秒；"
            + ("外联探针已阻断" if isolation.get("probe_blocked") is True else "外联探针未阻断"),
            str(isolation_status),
        )
    ]
    definitions = (
        ("离线端到端", "offline_e2e", "完整检索链路"),
        ("非环回连接", "non_loopback_connections", "候选进程连接采样"),
        ("路径穿越", "path_traversal", "授权根目录边界"),
        ("重解析点逃逸", "reparse_point_escape", "符号链接或目录联接"),
        ("正式包审计", "package_audit", "私有数据、凭据与绝对路径"),
    )
    rows.extend((label, detail, str(checks.get(key, "MISSING"))) for label, key, detail in definitions)
    return rows


def _build_performance_report(output: Path, context: dict[str, Any], comparison: dict[str, Any]) -> None:
    from docx import Document

    doc = Document()
    _style_document(doc, title="性能优化基准报告", subtitle="同机三轮基线、候选与准确率防回归", status=context["document_status"], commit=context["source_commit"])
    _add_gate_summary(doc, context)
    doc.add_heading("方法", level=1)
    workload = comparison.get("workload", {})
    doc.add_paragraph(
        "基线与候选使用同一硬件、电源模式、数据哈希、模型哈希和确定性工作负载。"
        "每轮使用与正式输入不重叠的预热输入不少于 10 次，再执行不少于 100 次混合冷热查询，共 3 轮，以轮间中位数比较。"
    )
    if isinstance(workload, dict):
        doc.add_paragraph(
            f"工作负载模式：{workload.get('workload_mode', '证据待生成')}；"
            f"唯一查询数：{workload.get('unique_queries', '证据待生成')}；"
            f"目标缓存命中率：{workload.get('target_cache_hit_ratio', '证据待生成')}；"
            f"预热输入与正式输入分离：{workload.get('warmup_inputs_disjoint', '证据待生成')}。"
        )
    doc.add_heading("比较结果", level=1)
    _add_table(doc, ("指标", "候选中位数", "相对基线变化"), _performance_summary_rows(comparison), [3600, 2880, 2880])
    doc.add_heading("准确率防回归", level=1)
    doc.add_paragraph("NQ 与 COCO 的 recall@10、MRR@10、NDCG@10 相对基线下降均不得超过 0.01。")
    doc.add_heading("结论", level=1)
    doc.add_paragraph(f"G6 当前状态：{context['gate_statuses'].get('G6', 'MISSING')}。热缓存嵌入、热缓存向量查询和峰值 RSS 改善均须达到 5%；混合负载与全链路 P95 不得回退超过 5%。")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def _build_security_report(output: Path, context: dict[str, Any], bugs: dict[str, Any], security: dict[str, Any]) -> None:
    from docx import Document

    doc = Document()
    _style_document(doc, title="缺陷修复与本地数据安全审查报告", subtitle="严重缺陷关闭、离线边界与包内容审计", status=context["document_status"], commit=context["source_commit"])
    _add_gate_summary(doc, context)
    doc.add_heading("缺陷严重度与关闭要求", level=1)
    doc.add_paragraph("Critical 包括数据损坏、越权读取、内容外发、无法启动或安全边界失效；High 包括核心索引/检索/打开/删除/持久化流程的稳定失败。最终要求 0 Open Critical、0 Open High。")
    _add_table(doc, ("统计项", "数量"), (
        ("Open Critical", bugs.get("open_critical", "证据待生成")),
        ("Open High", bugs.get("open_high", "证据待生成")),
        ("Closed", bugs.get("closed", "证据待生成")),
    ), [4800, 4560])
    doc.add_heading("离线与路径安全", level=1)
    _add_table(doc, ("检查", "范围或证据", "状态"), _security_summary_rows(security), [2600, 4760, 2000])
    isolation = security.get("network_isolation", {})
    if isinstance(isolation, dict):
        doc.add_paragraph(
            f"隔离方式为 {isolation.get('method', '证据待生成')}，连接审计持续 "
            f"{isolation.get('sample_seconds', '证据待生成')} 秒；最终门禁要求隔离已强制启用、外联探针被阻断，且全程无已建立的非环回连接。"
        )
    doc.add_heading("隐私结论", level=1)
    doc.add_paragraph("正式包不得包含用户索引、受控输入、日志、偏好、凭据、开发缓存或本机绝对工作路径；运行期间候选进程不得建立非环回连接。")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--coverage", type=Path)
    parser.add_argument("--stress", type=Path)
    parser.add_argument("--performance", type=Path)
    parser.add_argument("--bugs", type=Path)
    parser.add_argument("--security", type=Path)
    parser.add_argument("--draft", action="store_true")
    args = parser.parse_args()
    manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
    context = validate_report_inputs(manifest, draft=args.draft)
    args.output_dir.mkdir(parents=True, exist_ok=True)
    _build_test_report(args.output_dir / REPORT_FILENAMES[0], context, _read_optional(args.coverage), _read_optional(args.stress))
    _build_performance_report(args.output_dir / REPORT_FILENAMES[1], context, _read_optional(args.performance))
    _build_security_report(args.output_dir / REPORT_FILENAMES[2], context, _read_optional(args.bugs), _read_optional(args.security))
    for name in REPORT_FILENAMES:
        print(args.output_dir / name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
