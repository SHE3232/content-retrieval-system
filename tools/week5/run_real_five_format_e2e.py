#!/usr/bin/env python3
"""Run the Week 5 five-format workflow against a live MVP backend."""

from __future__ import annotations

import argparse
import json
import tempfile
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


TERMINAL = {"completed", "completed_with_errors", "failed"}


def request_json(base_url: str, method: str, path: str, payload=None):
    body = None
    headers = {"Accept": "application/json"}
    if payload is not None:
        body = json.dumps(payload).encode("utf-8")
        headers["Content-Type"] = "application/json"
    request = urllib.request.Request(
        urllib.parse.urljoin(base_url.rstrip("/") + "/", path.lstrip("/")),
        data=body,
        headers=headers,
        method=method,
    )
    try:
        with urllib.request.urlopen(request, timeout=120) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"{method} {path} -> HTTP {error.code}: {detail}") from error


def wait_for_job(base_url: str, job_id: str, timeout_seconds: int = 600):
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        value = request_json(base_url, "GET", f"/v1/indexing/jobs/{job_id}")
        if value["status"] in TERMINAL:
            return value
        time.sleep(1)
    raise TimeoutError(f"indexing job {job_id} did not finish in time")


def create_fixtures(root: Path):
    token = uuid.uuid4().hex[:10]
    text_tokens = {
        "week5-notes.txt": f"oriontext{token}",
        "week5-guide.pdf": f"nebula{token}",
        "week5-design.docx": f"quasar{token}",
    }
    (root / "week5-notes.txt").write_text(
        f"Week 5 controlled fixture. Unique term: {text_tokens['week5-notes.txt']}",
        encoding="utf-8",
    )

    pdf = canvas.Canvas(str(root / "week5-guide.pdf"), pagesize=A4)
    pdf.setFont("Helvetica", 18)
    pdf.drawString(72, 760, "Week 5 controlled PDF fixture")
    pdf.drawString(72, 720, text_tokens["week5-guide.pdf"])
    pdf.save()

    document = Document()
    document.add_heading("Week 5 controlled DOCX fixture", level=1)
    document.add_paragraph(text_tokens["week5-design.docx"])
    document.save(root / "week5-design.docx")

    jpeg = Image.new("RGB", (512, 512), "white")
    draw = ImageDraw.Draw(jpeg)
    draw.ellipse((96, 96, 416, 416), fill=(220, 30, 30), outline=(90, 0, 0), width=12)
    draw.polygon([(256, 105), (285, 45), (312, 112)], fill=(20, 130, 45))
    jpeg.save(root / "week5-red-apple.jpg", quality=95)

    png = Image.new("RGB", (512, 512), "white")
    draw = ImageDraw.Draw(png)
    draw.rectangle((96, 96, 416, 416), fill=(25, 90, 220), outline=(0, 25, 90), width=12)
    png.save(root / "week5-blue-square.png")
    return text_tokens


def search(base_url: str, query: str, channels: list[str]):
    return request_json(
        base_url,
        "POST",
        "/v1/search",
        {"query": query, "top_k": 10, "channels": channels, "filters": {}},
    )


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--base-url", default="http://127.0.0.1:8000")
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    fixture_root = Path(tempfile.gettempdir()) / f"contentretrieval-week5-{uuid.uuid4().hex}"
    fixture_root.mkdir(parents=True)
    tokens = create_fixtures(fixture_root)
    fixture_names = sorted(path.name for path in fixture_root.iterdir())
    report = {
        "status": "FAIL",
        "fixture_names": fixture_names,
        "indexing": {},
        "searches": {},
        "mutations": {},
    }
    source_keys = []
    try:
        created = request_json(
            args.base_url,
            "POST",
            "/v1/indexing/jobs",
            {
                "paths": [str(fixture_root)],
                "authorized_roots": [str(fixture_root)],
                "recursive": True,
            },
        )
        indexed = wait_for_job(args.base_url, created["job_id"])
        report["indexing"] = {"status": indexed["status"], "result": indexed.get("result")}
        if indexed["status"] != "completed" or indexed["result"]["indexed_files"] != 5:
            raise AssertionError(f"five files were not indexed successfully: {indexed}")

        catalog = request_json(args.base_url, "GET", "/v1/index/files?page=1&page_size=100")
        ours = [item for item in catalog["items"] if item["name"] in fixture_names]
        if len(ours) != 5:
            raise AssertionError(f"catalog contains {len(ours)} of five fixtures")
        source_keys = [item["source_key"] for item in ours]

        for name, token in tokens.items():
            response = search(args.base_url, token, ["keyword"])
            names = [hit["name"] for hit in response["hits"]]
            report["searches"][name] = {"query_kind": "keyword", "hit_names": names}
            if name not in names:
                raise AssertionError(f"keyword search did not return {name}: {names}")

        for name, query in (
            ("week5-red-apple.jpg", "a simple red apple on a white background"),
            ("week5-blue-square.png", "a simple blue square on a white background"),
        ):
            response = search(args.base_url, query, ["image_semantic"])
            names = [hit["name"] for hit in response["hits"]]
            report["searches"][name] = {"query_kind": "image_semantic", "hit_names": names}
            if not names or names[0] != name:
                raise AssertionError(f"image search did not rank {name} first: {names}")

        txt = next(item for item in ours if item["name"] == "week5-notes.txt")
        reindex = request_json(
            args.base_url,
            "POST",
            f"/v1/index/files/{txt['source_key']}/reindex",
        )
        reindexed = wait_for_job(args.base_url, reindex["job_id"])
        if reindexed["status"] != "completed":
            raise AssertionError(f"reindex did not complete: {reindexed}")
        report["mutations"]["reindex"] = "PASS"

        deleted = request_json(
            args.base_url,
            "DELETE",
            f"/v1/index/files/{txt['source_key']}",
        )
        if deleted["deleted_records"] < 1:
            raise AssertionError(f"remove mutation deleted no records: {deleted}")
        report["mutations"]["remove"] = "PASS"
        source_keys.remove(txt["source_key"])
        report["status"] = "PASS"
    except Exception as error:
        report["error"] = str(error)
        raise
    finally:
        for source_key in source_keys:
            try:
                request_json(args.base_url, "DELETE", f"/v1/index/files/{source_key}")
            except Exception as cleanup_error:
                report.setdefault("cleanup_errors", []).append(str(cleanup_error))
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
