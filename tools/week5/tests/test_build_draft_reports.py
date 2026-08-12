from pathlib import Path

from docx import Document

from tools.week5.build_draft_reports import main


def test_generated_reports_have_expected_structure(monkeypatch, tmp_path: Path):
    evidence = Path("docs/week5/evidence")
    monkeypatch.setattr(
        "sys.argv",
        ["build_draft_reports.py", str(evidence), str(tmp_path)],
    )

    main()

    expected = {
        "无障碍合规验证报告.docx": ("无障碍合规验证报告", "BLOCKED"),
        "UI可用性测试报告.docx": ("UI 可用性测试报告", "BLOCKED"),
        "无障碍用户指南（草稿）.docx": ("无障碍用户指南（草稿）", "草稿"),
    }
    for name, required in expected.items():
        path = tmp_path / name
        assert path.exists()
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for value in required:
            assert value in text
        section = document.sections[0]
        assert round(section.page_width.mm) == 210
        assert round(section.page_height.mm) == 297
        assert "C:\\Users" not in text
        assert "F:\\" not in text

    accessibility = Document(tmp_path / "无障碍合规验证报告.docx")
    accessibility_text = "\n".join(
        paragraph.text for paragraph in accessibility.paragraphs
    )
    assert "10/19 严格门禁 PASS" in accessibility_text
    assert "当前全量 Flutter 测试 178 项通过" in accessibility_text
    assert "Android release、Linux release、Web release 与 Windows release" in accessibility_text
    assert "WAVE 四状态扫描" in accessibility_text

    guide = Document(tmp_path / "无障碍用户指南（草稿）.docx")
    guide_text = "\n".join(paragraph.text for paragraph in guide.paragraphs)
    assert "Android release 已完成构建和设备启动" in guide_text
    assert "Accessibility Scanner 最终复扫按项目决定延后" in guide_text
    assert "Linux release 已完成 WSLg 启动检查" in guide_text
    assert "当前环境缺少 Android SDK" not in guide_text


def test_all_table_cells_use_white_fill(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "sys.argv",
        ["build_draft_reports.py", "docs/week5/evidence", str(tmp_path)],
    )
    main()

    for path in tmp_path.glob("*.docx"):
        document = Document(path)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    shading = cell._tc.get_or_add_tcPr().find(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
                    )
                    assert shading is not None
                    assert shading.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"
                    ) == "FFFFFF"
