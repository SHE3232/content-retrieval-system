#!/usr/bin/env python3
"""Build evidence-backed Week 5 draft DOCX reports."""

from __future__ import annotations

import argparse
from collections import Counter
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
REPORT_DATE = "2026-08-12"


def set_font(run, size=11, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def configure(doc: Document, running_title: str):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in (
        ("Title", 24, 0, 12),
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = name != "Title"
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.left_indent = Inches(0)
    header.paragraph_format.right_indent = Inches(0)
    header.paragraph_format.first_line_indent = Inches(0)
    set_font(header.add_run(running_title), 9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("第 "), 9)
    _field(footer, "PAGE")
    set_font(footer.add_run(" 页"), 9)
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()


def _field(paragraph, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for item in (begin, instr, separate, text, end):
        run.append(item)


def cover(doc: Document, title: str, subtitle: str, status: str, source_commit: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), 24, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), 14)
    p.paragraph_format.space_after = Pt(28)
    for label, value in (
        ("状态", status),
        ("版本", "Week 5 Draft 1.1"),
        ("日期", REPORT_DATE),
        ("应用源码提交", source_commit),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(f"{label}："), 11, bold=True)
        set_font(p.add_run(value), 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("本文件如实保留未执行验证项；不得作为 19/19 完成证明。"), 10, italic=True)
    doc.add_page_break()


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), 11, bold=True)
        set_font(p.add_run(text[len(bold_prefix) :]), 11)
    else:
        set_font(p.add_run(text), 11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(text), 11)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        set_font(cell.paragraphs[0].add_run(text), 9, bold=True)
    _repeat_header(tbl.rows[0])
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[index].text = ""
            set_font(cells[index].paragraphs[0].add_run(str(value)), 9)
    if widths:
        for row in tbl.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    for row in tbl.rows:
        _prevent_split(row)
        for cell in row.cells:
            _white_cell(cell)
    doc.add_paragraph()
    return tbl


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _white_cell(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), "FFFFFF")


def load_records(root: Path):
    manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))
    records = []
    for relative_path in manifest["records"]:
        path = root / relative_path
        value = json.loads(path.read_text(encoding="utf-8"))
        value["_path"] = path.relative_to(root).as_posix()
        records.append(value)
    return sorted(records, key=lambda value: value["gate_id"])


def application_source_commit(records):
    commits = [
        record["source_commit"]
        for record in records
        if record.get("status") == "PASS" and record.get("source_commit")
    ]
    if not commits:
        return "未记录"
    return Counter(commits).most_common(1)[0][0]


def build_accessibility(records, output):
    doc = Document()
    configure(doc, "Week 5 合规报告")
    passed = sum(record["status"] == "PASS" for record in records)
    source_commit = application_source_commit(records)
    cover(doc, "无障碍合规验证报告", "本地多模态内容检索系统", f"BLOCKED（{passed}/19 严格门禁 PASS）", source_commit)
    heading(doc, "1. 执行摘要")
    para(doc, f"严格完成状态为 BLOCKED。当前 19 项门禁中已有 {passed} 项通过，包括 Android release、Linux release、Web release 与 Windows release、真实五格式 E2E、设置重启持久化、高对比度、200% 字体和减少动态效果；其余项目保留为 BLOCKED。自动化无障碍测试证明实现行为，但不替代 NVDA、VoiceOver、Accessibility Scanner、WAVE 或人工走查。")
    heading(doc, "2. 范围与实现概况")
    for item in (
        "Flutter Material 3 搜索、索引库、设置三页已接入真实导航。",
        "已实现语义标题、状态实时播报、Ctrl/Cmd+1/2/3、Ctrl/Cmd+K、F5、纯键盘可操作基础。",
        "系统字号与用户字号相乘并封顶 200%；系统与用户减少动态效果采用逻辑或。",
        "明暗高对比度主题的关键颜色对已通过 WCAG 4.5:1/3:1 自动化比值检查。",
        "Android/Web 的目录选择和本地文件打开明确禁用；这两个 target 仅用于跨平台与指定工具验证。",
    ):
        bullet(doc, item)
    heading(doc, "3. 要求—证据矩阵")
    rows = []
    for record in records:
        issue = "；".join(record.get("issues", [])) or "无"
        rows.append((record["gate_id"], record["status"], record["observations"][0], issue))
    table(doc, ("门禁", "状态", "实际结果", "后续动作/问题"), rows, (1.25, 0.85, 2.2, 2.4))
    heading(doc, "4. 自动化 Flutter 无障碍检查")
    para(doc, "自动化套件覆盖语义标题与标签、点击目标规则、键盘快捷键、高对比度比值、200% 布局、设置持久化及状态播报。当前全量 Flutter 测试 177 项通过，静态分析无问题；后端 450 项通过、5 项跳过。该结论只描述代码层门禁。")
    heading(doc, "5. 平台辅助技术验证")
    for title, body in (
        ("Windows / NVDA", "未执行完整 NVDA 人工流程，状态 BLOCKED。"),
        ("macOS / VoiceOver", "无 macOS 环境，构建与 VoiceOver 均为 BLOCKED。"),
        ("Android / Accessibility Scanner", "Android release APK 已构建、签名校验并在 API 36 模拟器完成设置持久化复验；官方 Accessibility Scanner 尚未执行，工具门禁仍为 BLOCKED。"),
        ("Linux / WSLg", "Linux release 已完成构建、启动和可视化检查，状态 PASS。"),
        ("Web / WAVE", "Web release 已构建；尚未执行 WAVE 多状态审查，状态 BLOCKED。"),
    ):
        heading(doc, title, 2)
        para(doc, body)
    heading(doc, "6. 剩余风险与验收决定")
    para(doc, "不得将本报告作为第五周全部完成证明。完成条件是 evidence validator 在不带 --allow-incomplete 时输出 19/19，并且所有记录对应同一应用源码提交。")
    heading(doc, "7. 证据索引")
    table(doc, ("门禁", "记录路径", "附件数"), [(r["gate_id"], r["_path"], len(r["attachments"])) for r in records], (2.0, 3.8, 0.9))
    doc.save(output)


def build_usability(records, output):
    doc = Document()
    configure(doc, "Week 5 可用性报告")
    cover(doc, "UI 可用性测试报告", "本地多模态内容检索系统", "BLOCKED（0/3 参与者已完成）", application_source_commit(records))
    heading(doc, "1. 目标与工作流")
    para(doc, "目标是验证用户能否完成后端状态确认、搜索与筛选、结果打开/复制、索引库添加/进度/重建/移除以及无障碍设置。")
    heading(doc, "2. 参与者与隐私")
    para(doc, "P01、P02、P03 尚未招募或完成测试，因此不存在可报告的完成率、耗时、错误数、协助率、引语或 SUS 分数。后续会先获取同意，仅保留匿名标识和任务级数据。")
    table(doc, ("参与者", "状态", "输入方式", "结果"), (("P01", "BLOCKED", "待安排", "无数据"), ("P02", "BLOCKED", "待安排", "无数据"), ("P03", "BLOCKED", "至少一人纯键盘", "无数据")), (1.0, 1.0, 2.0, 2.7))
    heading(doc, "3. 预定测试脚本")
    for item in (
        "确认后端连接状态；离线时执行重新检测。",
        "输入查询，切换模式/内容类型，检查空、错误、加载和结果状态。",
        "打开和复制结果，确认错误反馈不泄露内部异常。",
        "选择目录、观察索引进度和失败详情，执行重新索引与从索引移除。",
        "切换主题、高对比度、200% 字号与减少动态效果，保存后重启验证。",
    ):
        bullet(doc, item)
    heading(doc, "4. 成功标准与记录字段")
    table(doc, ("字段", "记录规则"), (("任务成功", "独立完成/协助后完成/失败"), ("时间", "每项秒数；报告原始值与中位数"), ("错误", "误操作与恢复方式"), ("协助", "主持人提示次数"), ("发现", "严重度、页面、复现、建议")), (1.25, 5.45))
    heading(doc, "5. 当前可报告发现")
    para(doc, "代码和自动化审查已确认三页功能、响应状态和无障碍设置存在，但这些不是参与者可用性证据。当前唯一有效结论是可用性验收尚未完成。")
    heading(doc, "6. 后续执行与接受决定")
    para(doc, "完成三场主持式测试，其中至少一场全程纯键盘；汇总三个原始时间值、中位数、错误数和协助情况；修复重要问题后重测。完成前报告保持草稿和 BLOCKED。")
    doc.save(output)


def build_guide(records, output):
    doc = Document()
    configure(doc, "Week 5 用户指南（草稿）")
    cover(doc, "无障碍用户指南（草稿）", "本地多模态内容检索系统", "草稿；辅助技术跨平台验证待完成", application_source_commit(records))
    heading(doc, "1. 适用范围")
    para(doc, "Windows 是当前主要运行平台。Android release 已完成构建和设备启动，用于后续 Accessibility Scanner 验证；Linux release 已完成 WSLg 启动检查；Web release 已能构建，用于 WAVE 验证。macOS 构建与 VoiceOver 仍待真实 Mac 环境完成。")
    heading(doc, "2. 启动应用")
    for item in (
        "在项目根目录运行 start-mvp.ps1 -CheckOnly，确认输出 MVP preflight passed。",
        "启动后端并确认 /health/ready 返回 ready，再运行 Windows Flutter 应用。",
        "默认后端地址为 http://127.0.0.1:8000，可在“设置”中修改为无路径、无凭据、无查询参数的 HTTP(S) 根地址。",
        "模型文件已经存在且清单校验通过时不会在每次启动时重新下载；只有模型缺失或恢复流程明确要求时才需要下载。",
    ):
        bullet(doc, item)
    heading(doc, "3. 搜索与筛选")
    para(doc, "进入“搜索”，在“搜索内容”输入关键词并选择“搜索”。可使用检索模式、内容类型和检索通道筛选；至少保留一个检索通道。加载、空结果和错误均以文字状态显示并提供可执行恢复动作。")
    heading(doc, "4. 读取与处理结果")
    para(doc, "结果显示文件名、匹配摘要和元数据。选择“打开”调用系统关联程序；选择“复制路径”将路径复制到剪贴板。Web/Android 不支持访问桌面后端的本地路径，因此相关操作会禁用或提示使用桌面版。")
    heading(doc, "5. 管理索引库")
    para(doc, "进入“索引库”，可刷新、添加文件夹、查看索引任务和失败详情。重新索引会更新搜索索引；“从索引移除”只删除索引记录，不删除磁盘原文件。所有破坏性动作都要求命名确认。")
    heading(doc, "6. 外观与无障碍设置")
    for item in (
        "主题：跟随系统、浅色或深色。",
        "高对比度：增强文字、边界和焦点区分。",
        "文字大小：100%、125%、150% 或 200%；与系统倍率组合后最高 200%。",
        "减少动态效果：与系统设置共同生效；任一开启即减少非必要动画。",
        "保存设置后写入本机 JSON 文件；文件损坏时恢复安全默认值并显示提示。",
    ):
        bullet(doc, item)
    doc.add_page_break()
    heading(doc, "7. 键盘快捷键")
    table(doc, ("快捷键", "作用"), (("Ctrl/Cmd+K", "聚焦搜索框"), ("Ctrl/Cmd+1", "打开搜索"), ("Ctrl/Cmd+2", "打开索引库"), ("Ctrl/Cmd+3", "打开设置"), ("F5", "刷新当前网络页面"), ("Tab / Shift+Tab", "向前/向后移动焦点"), ("Enter / Space", "激活当前控件"), ("Escape", "关闭临时界面或取消搜索框焦点")), (2.0, 4.7))
    heading(doc, "8. 屏幕阅读器快速开始")
    heading(doc, "NVDA（Windows）", 2)
    para(doc, "启动 NVDA 后使用 Tab 导航三项主导航和页面控件。状态变化会通过实时区域播报。完整 NVDA 验收尚未执行；若发现重复标签、焦点丢失或状态未播报，请记录页面、操作和实际播报。")
    heading(doc, "VoiceOver（macOS）", 2)
    para(doc, "使用 VoiceOver 键盘命令遍历导航、表单、结果和对话框。当前无 macOS 实测记录，因此本节为待验证说明。")
    heading(doc, "9. 常见问题")
    table(doc, ("现象", "处理"), (("后端离线", "确认服务和地址，选择“重新检测”"), ("模型清单缺失", "恢复 models/model-manifest.json 及对应权重，再运行预检"), ("Android Scanner 未执行", "在 Google Play 模拟器安装官方 Accessibility Scanner，启用服务后记录完整工作流"), ("索引任务失败", "打开失败详情，确认路径权限和文件格式后重试"), ("设置数据损坏", "应用恢复默认值；重新设置并保存")), (2.1, 4.6))
    heading(doc, "10. 隐私与反馈")
    para(doc, "检索与设置以本机处理为主。不要在测试证据中提交私人文档内容、完整敏感路径或参与者身份。反馈应包含平台、应用提交、操作步骤、期望与实际结果。")
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("evidence", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)
    records = load_records(args.evidence)
    build_accessibility(records, args.output / "无障碍合规验证报告.docx")
    build_usability(records, args.output / "UI可用性测试报告.docx")
    build_guide(records, args.output / "无障碍用户指南（草稿）.docx")


if __name__ == "__main__":
    main()
