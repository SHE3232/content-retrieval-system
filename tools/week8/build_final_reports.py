#!/usr/bin/env python3
"""Generate the evidence-bound Week 8 final report and weekly report DOCX files."""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

FONT_NAME = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
WHITE_HEX = "FFFFFF"
BLACK_HEX = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
REQUIRED_CHAPTERS = (
    "项目背景与目标",
    "需求分析与验收口径",
    "总体架构与技术选型",
    "八周实施过程与阶段成果",
    "文件解析与统一内容模型",
    "多模态嵌入与模型工程",
    "向量存储、索引一致性与生命周期",
    "关键词、语义与混合检索",
    "Flutter 客户端与交互设计",
    "无障碍设计与验证",
    "系统集成、性能与稳定性",
    "本地数据安全、隐私与开源合规",
    "测试体系、指标与最终验收",
    "关键问题、解决过程与工程反思",
    "最终交付、局限与后续规划",
)


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.color.rgb = BLACK
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), BLACK_HEX)


def _set_style_font(style: Any, *, size: float, bold: bool = False) -> None:
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), BLACK_HEX)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Title": (26, True, 0, 8),
        "Subtitle": (14, False, 0, 10),
        "Heading 1": (16, True, 18, 10),
        "Heading 2": (13, True, 12, 6),
        "Heading 3": (12, True, 8, 4),
        "Caption": (10, False, 4, 10),
        "Header": (9, False, 0, 0),
        "Footer": (9, False, 0, 0),
    }
    for name, (size, bold, before, after) in heading_tokens.items():
        style = styles[name]
        _set_style_font(style, size=size, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading") or name == "Caption"
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _configure_section(section: Any) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def _append_field(paragraph: Any, instruction: str, display: str) -> None:
    run = paragraph.add_run()
    _set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def _configure_header_footer(document: Document) -> None:
    for section in document.sections:
        _configure_section(section)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.text = "离线可访问多模态本地内容检索系统｜项目结项"
        for run in header.runs:
            _set_run_font(run, size=9)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix = footer.add_run("第 ")
        _set_run_font(prefix, size=9)
        _append_field(footer, "PAGE", "1")
        suffix = footer.add_run(" 页")
        _set_run_font(suffix, size=9)


def _shade_cell(cell: Any, fill: str = WHITE_HEX) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", 100), ("bottom", 100), ("start", 120), ("end", 120)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BLACK_HEX)


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("table column widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _shade_cell(cell)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_table_borders(table)


def _format_table_text(table: Any) -> None:
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            marker = OxmlElement("w:tblHeader")
            marker.set(qn("w:val"), "true")
            row._tr.get_or_add_trPr().append(marker)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    _set_run_font(run, size=9.5, bold=row_index == 0)


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    caption: str,
) -> None:
    caption_paragraph = document.add_paragraph(caption, style="Caption")
    caption_paragraph.paragraph_format.keep_with_next = True
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    _set_table_geometry(table, widths_dxa)
    _format_table_text(table)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def _test_summary(evidence: dict[str, Any]) -> str:
    tests = evidence.get("tests", {})
    if not isinstance(tests, dict):
        return "没有可用的测试套件记录"
    parts = []
    for name, result in tests.items():
        if isinstance(result, dict):
            parts.append(
                f"{name}：{result.get('status', 'BLOCKED')}，"
                f"通过 {int(result.get('passed', 0))}，跳过 {int(result.get('skipped', 0))}"
            )
    return "；".join(parts)


def _platform_status(evidence: dict[str, Any], platform: str) -> str:
    platforms = evidence.get("platforms", {})
    result = platforms.get(platform, {}) if isinstance(platforms, dict) else {}
    if not isinstance(result, dict):
        return "BLOCKED"
    status = str(result.get("status", "BLOCKED"))
    reason = str(result.get("reason", "")).strip()
    return f"{status}（{reason}）" if reason else status


def _replacements(evidence: dict[str, Any]) -> dict[str, str]:
    benchmarks = evidence.get("benchmarks", {})
    if not isinstance(benchmarks, dict):
        benchmarks = {}
    return {
        "{{SOURCE_COMMIT}}": str(evidence["source_commit"]),
        "{{TEST_SUMMARY}}": _test_summary(evidence),
        "{{WINDOWS_STATUS}}": _platform_status(evidence, "windows"),
        "{{LINUX_STATUS}}": _platform_status(evidence, "linux"),
        "{{MACOS_STATUS}}": _platform_status(evidence, "macos"),
        "{{SEARCH_P95_MS}}": f"{float(benchmarks.get('search_p95_ms', 0)):.2f}",
        "{{TARGET_P95_MS}}": f"{float(benchmarks.get('target_p95_ms', 0)):.2f}",
        "{{TEXT_BATCH1_P50_MS}}": f"{float(benchmarks.get('text_batch1_p50_ms', 0)):.2f}",
        "{{TEXT_BATCH16_THROUGHPUT}}": f"{float(benchmarks.get('text_batch16_throughput', 0)):.2f}",
    }


def _replace(text: str, replacements: dict[str, str]) -> str:
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


def _table_data(name: str, evidence: dict[str, Any]) -> tuple[list[str], list[list[str]], list[int], str]:
    if name == "acceptance":
        return (
            ["维度", "验收内容", "证据规则"],
            [
                ["功能", "五格式摄取、关键词/语义检索、索引管理", "自动化与五格式端到端"],
                ["性能", "记录 P50、P95、吞吐和硬件条件", "只引用结构化基线"],
                ["无障碍", "键盘、语义、高对比度、文本缩放、减少动效", "按平台与工具分开判定"],
                ["发布", "目标平台 Release、许可证、版本、哈希", "直接证据才可 PASS"],
            ],
            [1800, 3960, 3600],
            "表 1 需求与验收证据矩阵",
        )
    if name == "models":
        return (
            ["模型", "空间/维度", "默认发行", "许可证边界"],
            [
                ["多语言文本模型", "text-semantic-v1 / 384", "包含", "Apache-2.0"],
                ["MobileCLIP-S0", "mobileclip-image-text-v1 / 512", "不包含", "权重仅限非商业研究"],
            ],
            [2300, 2600, 1700, 2760],
            "表 2 模型、向量空间与发行边界",
        )
    if name == "week_timeline":
        rows = [
            ["1", "需求、架构与验收口径", "PRODUCT、DESIGN、架构基线"],
            ["2", "五格式发现与解析", "解析器注册表、统一内容模型"],
            ["3", "文本与图文嵌入", "模型清单、NQ/COCO 与性能证据"],
            ["4", "Chroma 索引与混合检索", "FastAPI MVP、RRF、生命周期"],
            ["5", "Flutter 与无障碍", "桌面客户端、五格式 E2E"],
            ["6", "集成、性能与打包", "启动器、故障注入、发布门禁"],
            ["7", "文档、演示与对账", "九段式脚本、证据索引"],
            ["8", "清理、平台候选与结项", "干净源码、报告、统一清单"],
        ]
        return ["周次", "工作重点", "核心产出"], rows, [900, 3500, 4960], "表 3 八周实施与产出"
    if name == "parsers":
        return (
            ["格式", "解析路径", "关键边界"],
            [
                ["TXT", "本地字节解码", "UTF-8、带 BOM 的 UTF-16/32；失败关闭"],
                ["PDF", "本地 Apache Tika", "超时、空正文和服务不可用分开记录"],
                ["DOCX", "本地 Apache Tika", "不执行文档宏；只提取内容"],
                ["JPEG / PNG", "Pillow", "校验格式、尺寸与颜色模式"],
            ],
            [1500, 2600, 5260],
            "表 4 五格式解析策略",
        )
    if name == "retrieval":
        return (
            ["通道", "适用问题", "排序与边界"],
            [
                ["关键词", "术语、编号、文件名", "字段加权 BM25"],
                ["文本语义", "措辞不同但语义接近", "384 维余弦空间"],
                ["图文语义", "自然语言查找图片", "512 维研究配置"],
                ["混合", "综合精确匹配与语义召回", "加权 RRF 文件级融合"],
            ],
            [1800, 3600, 3960],
            "表 5 检索通道与融合规则",
        )
    if name == "accessibility":
        return (
            ["能力", "实现", "验收"],
            [
                ["键盘", "焦点顺序与 Enter/Space 等价操作", "组件测试与 Windows 记录"],
                ["语义", "标签、提示、状态播报", "语义树自动化"],
                ["视觉", "高对比度、200% 文本、可换行布局", "多状态截图与测试"],
                ["动态", "减少或取消非必要过渡", "偏好持久化与组件测试"],
                ["辅助技术", "NVDA / VoiceOver / Scanner", "无直接环境时保持 BLOCKED"],
            ],
            [1700, 4300, 3360],
            "表 6 无障碍能力与证据",
        )
    if name == "test_summary":
        tests = evidence.get("tests", {})
        rows = []
        if isinstance(tests, dict):
            for suite, result in tests.items():
                if isinstance(result, dict):
                    rows.append(
                        [
                            str(suite),
                            str(result.get("status", "BLOCKED")),
                            str(int(result.get("passed", 0))),
                            str(int(result.get("skipped", 0))),
                        ]
                    )
        return ["套件", "状态", "通过", "跳过"], rows, [4200, 1800, 1680, 1680], "表 7 最终测试套件结果"
    if name == "delivery_status":
        external = evidence.get("external_gates", {})
        if not isinstance(external, dict):
            external = {}
        rows = [
            ["Windows", _platform_status(evidence, "windows")],
            ["Linux", _platform_status(evidence, "linux")],
            ["macOS", _platform_status(evidence, "macos")],
            ["GitHub 公开发布", str(external.get("github", "BLOCKED"))],
            ["真实五分钟演示视频", str(external.get("video", "BLOCKED"))],
        ]
        return ["交付门禁", "最终状态与说明"], rows, [2800, 6560], "表 8 最终交付状态"
    raise ValueError(f"unknown report table marker: {name}")


def _add_picture(document: Document, path: Path, caption: str) -> None:
    if not path.is_file():
        raise ValueError(f"report image is missing: {path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(6.25))
    inline_shape._inline.docPr.set("descr", caption)
    caption_paragraph = document.add_paragraph(caption, style="Caption")
    caption_paragraph.paragraph_format.keep_together = True


def _add_cover(document: Document, evidence: dict[str, Any], *, weekly: bool = False) -> None:
    for _ in range(5 if not weekly else 2):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("WEEK 8｜FINAL DELIVERY" if not weekly else "WEEK 8｜WEEKLY REPORT")
    _set_run_font(run, size=12, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("离线可访问多模态本地内容检索系统")
    _set_run_font(title_run, size=26, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("第八周工作周报" if weekly else "项目结项报告")
    _set_run_font(subtitle_run, size=18, bold=True)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    meta_text = (
        f"源码提交：{evidence['source_commit']}\n"
        f"证据时间：{evidence.get('generated_at', '')}\n"
        "设计基线：narrative_proposal；命名覆盖 formal_monochrome_a4"
    )
    meta_run = meta.add_run(meta_text)
    _set_run_font(meta_run, size=10)
    document.add_page_break()


def _add_static_toc(document: Document) -> None:
    heading = document.add_paragraph("目录", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, chapter in enumerate(REQUIRED_CHAPTERS, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{index}. {chapter}")
        _set_run_font(run, size=11)
    document.add_page_break()


def _build_final_report(
    evidence: dict[str, Any],
    content_path: Path,
    assets_dir: Path,
    screenshot_dir: Path,
    output: Path,
) -> None:
    source = content_path.read_text(encoding="utf-8")
    replacements = _replacements(evidence)
    source = _replace(source, replacements)
    if re.search(r"\{\{[^}]+\}\}", source):
        raise ValueError("unresolved report content token")

    document = Document()
    _configure_styles(document)
    _configure_section(document.sections[0])
    _add_cover(document, evidence)
    _add_static_toc(document)

    for raw_line in source.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_paragraph(line[2:].strip(), style="Heading 1")
            continue
        if line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 2")
            continue
        figure = re.fullmatch(r"\[\[FIGURE:([^|]+)\|(.+)\]\]", line)
        if figure:
            _add_picture(document, assets_dir / figure.group(1), figure.group(2))
            continue
        screenshot = re.fullmatch(r"\[\[SCREENSHOT:([^|]+)\|(.+)\]\]", line)
        if screenshot:
            _add_picture(document, screenshot_dir / screenshot.group(1), screenshot.group(2))
            continue
        table_marker = re.fullmatch(r"\[\[TABLE:([^]]+)\]\]", line)
        if table_marker:
            _add_table(document, *_table_data(table_marker.group(1), evidence))
            continue
        paragraph = document.add_paragraph(line, style="Normal")
        paragraph.paragraph_format.widow_control = True

    _configure_header_footer(document)
    document.core_properties.title = "离线可访问多模态本地内容检索系统项目结项报告"
    document.core_properties.subject = "八周项目成果、验证证据与最终交付"
    document.core_properties.author = "Offline Accessible Multimodal Retrieval Project"
    document.core_properties.keywords = f"source_commit={evidence['source_commit']};week8;final-report"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def _build_weekly_report(evidence: dict[str, Any], assets_dir: Path, output: Path) -> None:
    document = Document()
    _configure_styles(document)
    _configure_section(document.sections[0])
    _add_cover(document, evidence, weekly=True)

    document.add_paragraph("一、本周目标与范围", style="Heading 1")
    document.add_paragraph(
        "第八周围绕三项总目标收口：完成平台发布、公开源码、演示和作品集待办；生成不少于八千字的图文结项报告；通过静态分析、引用搜索、测试和真实运行整理独立干净工程目录。全部产物以同一源码提交和统一清单为事实锚点。"
    )
    document.add_paragraph("二、已完成的本地工作", style="Heading 1")
    document.add_paragraph(
        "已完成公开依赖解耦、文本版能力降级、源代码白名单、死代码审计、Windows 与 Linux 构建脚本、macOS 真实主机门禁、社区治理文件、CI 与标签发布工作流、报告生成与交付验证工具。受限模型只进入独立研究包，默认公开包不含 MobileCLIP 权重。"
    )
    _add_table(document, *_table_data("test_summary", evidence))
    _add_picture(document, assets_dir / "06_测试结果汇总.png", "图 1 第八周测试套件汇总")

    document.add_paragraph("三、平台与外部门禁", style="Heading 1")
    external = evidence.get("external_gates", {})
    if not isinstance(external, dict):
        external = {}
    status_rows = [
        ["Windows", _platform_status(evidence, "windows")],
        ["Linux", _platform_status(evidence, "linux")],
        ["macOS", _platform_status(evidence, "macos")],
        ["GitHub", str(external.get("github", "BLOCKED：未配置可认证远程仓库"))],
        ["视频", str(external.get("video", "BLOCKED：需要真实人工操作与录音证据"))],
    ]
    _add_table(document, ["门禁", "状态"], status_rows, [2200, 7160], "表 2 平台与外部门禁")

    document.add_paragraph("四、清理与工程质量", style="Heading 1")
    document.add_paragraph(
        "清理过程保留原始仓库与 Git 历史，只在独立工作树修改。通过 Vulture、rg 引用搜索、Git 跟踪清单与回归测试确认删除对象；FastAPI 路由和校验器等装饰器入口保留为受审豁免。公开 Python 锁文件不再引用本机 MobileCLIP 源码，生成目录、模型、缓存、数据库、日志和用户数据均不进入公开工程。"
    )
    document.add_paragraph("五、风险与下一步", style="Heading 1")
    document.add_paragraph(
        "风险集中在真实 macOS 环境、公开 GitHub 权限和真实五分钟视频三项外部条件。下一步动作分别是：在 Darwin 主机运行发布脚本并完成 VoiceOver；获得远程仓库权限后推送冻结提交、等待 CI、创建 v1.0.0 与 Release；按九段式脚本完成两轮真实预演和最终录制。没有直接证据前这些状态保持 BLOCKED。"
    )
    document.add_paragraph("六、本周结论", style="Heading 1")
    document.add_paragraph(
        f"本地可执行工作已按证据推进，当前源码提交为 {evidence['source_commit']}。报告、平台候选和统一交付清单只引用该提交；任何后续跟踪文件变化都会触发重新冻结与重建，避免旧归档被手工修改提交号。"
    )

    _configure_header_footer(document)
    document.core_properties.title = "第八周工作周报"
    document.core_properties.author = "Offline Accessible Multimodal Retrieval Project"
    document.core_properties.keywords = f"source_commit={evidence['source_commit']};week8;weekly-report"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_reports(
    *,
    evidence_path: Path,
    content_path: Path,
    assets_dir: Path,
    screenshot_dir: Path,
    output_dir: Path,
) -> dict[str, object]:
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    commit = evidence.get("source_commit")
    if not isinstance(commit, str) or not re.fullmatch(r"[0-9a-f]{40}", commit):
        raise ValueError("report evidence requires a full lowercase source_commit")
    final_report = output_dir / "项目结项报告.docx"
    weekly_report = output_dir / "第八周工作周报.docx"
    _build_final_report(evidence, content_path, assets_dir, screenshot_dir, final_report)
    _build_weekly_report(evidence, assets_dir, weekly_report)
    result: dict[str, object] = {
        "schema_version": 1,
        "source_commit": commit,
        "generated_at": datetime.now().astimezone().isoformat(),
        "preset": "narrative_proposal",
        "named_override": "formal_monochrome_a4",
        "final_report": str(final_report.resolve()),
        "weekly_report": str(weekly_report.resolve()),
    }
    (output_dir / "report-build.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--evidence", type=Path, required=True)
    parser.add_argument("--content", type=Path, required=True)
    parser.add_argument("--assets", type=Path, required=True)
    parser.add_argument("--screenshots", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args(argv)
    result = build_reports(
        evidence_path=args.evidence,
        content_path=args.content,
        assets_dir=args.assets,
        screenshot_dir=args.screenshots,
        output_dir=args.output,
    )
    print(json.dumps(result, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
