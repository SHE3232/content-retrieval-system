from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from tools.week8.build_final_reports import REQUIRED_CHAPTERS, build_reports
from tools.week8.build_report_figures import build_figures

REPOSITORY = Path(__file__).resolve().parents[3]
CONTENT = REPOSITORY / "docs" / "week8" / "reports" / "项目结项报告正文.md"


def _evidence(path: Path) -> Path:
    path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "source_commit": "e" * 40,
                "generated_at": "2026-08-28T01:00:00+08:00",
                "tests": {
                    "backend": {"status": "PASS", "passed": 445, "skipped": 1},
                    "week8": {"status": "PASS", "passed": 51, "skipped": 1},
                    "flutter": {"status": "PASS", "passed": 249, "skipped": 0},
                },
                "platforms": {
                    "windows": {"status": "PASS", "reason": "实机构建与归档验证通过"},
                    "linux": {"status": "PASS", "reason": "Ubuntu 24.04 构建与归档验证通过"},
                    "macos": {"status": "BLOCKED", "reason": "缺少真实 macOS 主机"},
                },
                "benchmarks": {
                    "search_p95_ms": 239.292845017917,
                    "target_p95_ms": 2000.0,
                    "text_batch1_p50_ms": 21.179750037845224,
                    "text_batch16_throughput": 346.5024477276606,
                },
                "five_formats": {
                    "status": "PASS",
                    "parsed_files": 5,
                    "indexed_files": 5,
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


def test_generated_final_report_meets_content_and_structure_gates(
    tmp_path: Path,
) -> None:
    evidence = _evidence(tmp_path / "evidence.json")
    assets = tmp_path / "assets"
    build_figures(evidence, assets)
    output = tmp_path / "reports"

    result = build_reports(
        evidence_path=evidence,
        content_path=CONTENT,
        assets_dir=assets,
        screenshot_dir=REPOSITORY / "docs" / "week5" / "evidence" / "attachments",
        output_dir=output,
    )

    final_report = Path(result["final_report"])
    weekly_report = Path(result["weekly_report"])
    assert final_report.is_file()
    assert weekly_report.is_file()

    document = Document(final_report)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    body = "".join(
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Normal"
    )
    chinese_count = len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff]", body))

    assert chinese_count >= 8_000
    assert all(chapter in text for chapter in REQUIRED_CHAPTERS)
    assert len(document.inline_shapes) >= 12
    assert len(document.tables) >= 8
    assert "e" * 40 in text
    assert not re.search(r"TBD|TODO|\{\{[^}]+\}\}|占位符", text, re.IGNORECASE)

    with ZipFile(final_report) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        styles_xml = package.read("word/styles.xml").decode("utf-8")
        assert "Times New Roman" in document_xml + styles_xml
        assert 'w:w="11906"' in document_xml
        assert 'w:h="16838"' in document_xml
        assert 'w:fill="FFFFFF"' in document_xml
        assert 'w:color="000000"' in document_xml

    weekly_text = "\n".join(
        paragraph.text for paragraph in Document(weekly_report).paragraphs
    )
    assert "第八周工作周报" in weekly_text
    assert "macOS" in weekly_text and "BLOCKED" in weekly_text
    assert "GitHub" in weekly_text and "BLOCKED" in weekly_text
    assert "视频" in weekly_text and "BLOCKED" in weekly_text
