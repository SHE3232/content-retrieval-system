"""Create deterministic five-format project demonstration fixtures."""

import argparse
import copy
import json
import os
import shutil
import tempfile
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from PIL import Image, ImageDraw
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

EXPECTED_FILES = (
    "01_课程检索笔记.txt",
    "02_无障碍设计指南.pdf",
    "03_离线系统方案.docx",
    "04_红色苹果.jpg",
    "05_蓝色方块.png",
)

def _expected_entries() -> list[dict]:
    values = [("星桥检索协议", "精确"), ("哪个文档介绍了不用鼠标操作界面", "文本语义"),
              ("怎样在断网时保护本地文档隐私", "文本语义"),
              ("a simple red apple on a white background", "图像语义"),
              ("a simple blue square on a white background", "图像语义")]
    return [{"name": name, "query": query, "mode": mode} for name, (query, mode) in zip(EXPECTED_FILES, values)]


def _expected_manifest() -> dict:
    return {"schema_version": 1, "generated_by": "tools/demo/generate_demo_data.py", "files": _expected_entries()}


def _write_txt(path: Path) -> None:
    path.write_text(
        "课程资料整理说明\n\n本课介绍星桥检索协议。无需记住文件名，也能按内容检索相关资料，按内容找文件。\n",
        encoding="utf-8",
    )


def _write_pdf(path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    pdf = canvas.Canvas(str(path), pagesize=(595, 842))
    pdf.setFont("STSong-Light", 20)
    pdf.drawString(72, 770, "无障碍设计指南")
    pdf.setFont("STSong-Light", 12)
    lines = [
        "DEMO-PDF-ACCESSIBILITY",
        "支持键盘 Tab 完成界面导航。",
        "提供高对比度配色，字号支持 200% 放大。",
        "提供减少动态效果选项，降低视觉干扰。",
    ]
    for index, line in enumerate(lines):
        pdf.drawString(72, 730 - index * 28, line)
    pdf.save()


def _write_docx(path: Path) -> None:
    doc = Document()
    for style_name in ("Normal", "Title", "Heading 1"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_heading("离线系统方案", level=1)
    doc.add_paragraph(
        "系统在无网络时仍可运行。解析、检索与排序均在本机完成；用户文件、查询和片段不会发送到云端。"
    )
    doc.save(path)


def _write_images(directory: Path) -> None:
    apple = Image.new("RGB", (512, 512), "white")
    draw = ImageDraw.Draw(apple)
    draw.ellipse((112, 115, 402, 425), fill=(220, 30, 35), outline=(120, 0, 0), width=12)
    draw.polygon([(275, 135), (330, 72), (355, 78), (315, 150)], fill=(40, 145, 55))
    apple.save(directory / EXPECTED_FILES[3], quality=95)
    square = Image.new("RGB", (512, 512), "white")
    ImageDraw.Draw(square).rectangle((110, 110, 402, 402), fill=(35, 90, 220))
    square.save(directory / EXPECTED_FILES[4])


def _is_owned_manifest(path: Path) -> bool:
    try:
        manifest = json.loads(path.read_text(encoding="utf-8"))
        return manifest == _expected_manifest()
    except (OSError, UnicodeError, json.JSONDecodeError, KeyError, TypeError):
        return False


def _generate_into(directory: Path) -> None:
    _write_txt(directory / EXPECTED_FILES[0])
    _write_pdf(directory / EXPECTED_FILES[1])
    _write_docx(directory / EXPECTED_FILES[2])
    _write_images(directory)
    (directory / "MANIFEST.json").write_text(json.dumps(_expected_manifest(), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _publish(staging: Path, root: Path) -> None:
    rollback = root.parent / f".{root.name}.rollback-{uuid.uuid4().hex}"
    rollback.mkdir()
    known = (*EXPECTED_FILES, "MANIFEST.json")
    moved: list[str] = []
    published: list[str] = []
    try:
        for name in known:
            target = root / name
            if target.exists() and target.is_dir():
                raise FileExistsError(f"Known target is a directory: {target}")
        for name in known:
            target = root / name
            if target.exists() or target.is_symlink():
                os.replace(target, rollback / name); moved.append(name)
        for name in EXPECTED_FILES:
            os.replace(staging / name, root / name); published.append(name)
        os.replace(staging / "MANIFEST.json", root / "MANIFEST.json"); published.append("MANIFEST.json")
    except Exception:
        for name in reversed(published):
            target = root / name
            if target.exists() or target.is_symlink(): target.unlink()
        for name in reversed(moved):
            os.replace(rollback / name, root / name)
        raise
    finally:
        shutil.rmtree(rollback, ignore_errors=True)


def generate_demo_data(output: str | Path, force: bool = False) -> dict:
    """Generate five fixed demo files and return the manifest dictionary."""
    directory = Path(output).expanduser()
    directory.mkdir(parents=True, exist_ok=True)
    existing = list(directory.iterdir())
    manifest_path = directory / "MANIFEST.json"
    if existing and (not force or not _is_owned_manifest(manifest_path)):
        raise FileExistsError(f"Output directory is not empty: {directory}")
    staging = Path(tempfile.mkdtemp(prefix=f".{directory.name}.staging-", dir=directory.parent))
    try:
        _generate_into(staging)
        _publish(staging, directory)
    finally:
        shutil.rmtree(staging, ignore_errors=True)
    return copy.deepcopy(_expected_manifest())


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", type=Path)
    parser.add_argument("--force", action="store_true")
    args = parser.parse_args(argv)
    generate_demo_data(args.output, force=args.force)
    print(f"Generated five demo files in {args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
