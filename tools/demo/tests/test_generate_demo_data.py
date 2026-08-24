import json
import os
import sys
import subprocess
try:
    import tomllib
except ModuleNotFoundError:  # Python 3.10
    import tomli as tomllib
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile
from unittest import mock

from docx import Document
from PIL import Image
from pypdfium2 import PdfDocument

from tools.demo.generate_demo_data import EXPECTED_FILES, generate_demo_data


class DemoDataGeneratorTests(unittest.TestCase):
    def test_generates_exact_five_files_and_matching_manifest(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "demo"
            result = generate_demo_data(out)
            self.assertEqual(set(p.name for p in out.iterdir()), set(EXPECTED_FILES) | {"MANIFEST.json"})
            manifest = json.loads((out / "MANIFEST.json").read_text(encoding="utf-8"))
            self.assertEqual(manifest["schema_version"], 1)
            self.assertEqual(manifest["generated_by"], "tools/demo/generate_demo_data.py")
            self.assertEqual(len(manifest["files"]), 5)
            self.assertEqual(result, manifest)
            self.assertEqual({entry["name"] for entry in manifest["files"]}, set(EXPECTED_FILES))
            expected = [
                (EXPECTED_FILES[0], "星桥检索协议", "精确"),
                (EXPECTED_FILES[1], "哪个文档介绍了不用鼠标操作界面", "文本语义"),
                (EXPECTED_FILES[2], "怎样在断网时保护本地文档隐私", "文本语义"),
                (EXPECTED_FILES[3], "a simple red apple on a white background", "图像语义"),
                (EXPECTED_FILES[4], "a simple blue square on a white background", "图像语义"),
            ]
            self.assertEqual([(e["name"], e["query"], e["mode"]) for e in manifest["files"]], expected)
            for entry in manifest["files"]:
                self.assertIn("query", entry)
                self.assertIn("mode", entry)

    def test_text_and_office_documents_contain_demo_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            generate_demo_data(out)
            txt = (out / "01_课程检索笔记.txt").read_text(encoding="utf-8")
            self.assertIn("星桥检索协议", txt)
            self.assertIn("按内容检索", txt)
            self.assertIn("无需记住文件名", txt)
            doc = Document(out / "03_离线系统方案.docx")
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("无网络时仍可运行", text)
            self.assertIn("本机完成", text); self.assertIn("解析、检索与排序", text)
            self.assertIn("不会发送到云端", text)
            self.assertEqual(doc.styles["Normal"].font.name, "Times New Roman")
            self.assertEqual(doc.styles["Title"].font.name, "Times New Roman")
            self.assertEqual(doc.styles["Heading 1"].font.name, "Times New Roman")
            self.assertEqual(doc.styles["Heading 1"].font.color.rgb, (0, 0, 0))
            with ZipFile(out / "03_离线系统方案.docx") as archive:
                styles_xml = archive.read("word/styles.xml").decode("utf-8")
            self.assertIn('w:eastAsia="Times New Roman"', styles_xml)

    def test_pdf_and_images_are_real_and_contentful(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            generate_demo_data(out)
            pdf = PdfDocument(str(out / "02_无障碍设计指南.pdf"))
            extracted = "".join(page.get_textpage().get_text_bounded() for page in pdf)
            pdf.close()
            self.assertIn(b"STSong-Light", (out / "02_无障碍设计指南.pdf").read_bytes())
            for phrase in ("DEMO-PDF-ACCESSIBILITY", "Tab", "高对比度", "200%", "减少动态效果"):
                self.assertIn(phrase, extracted)
            with Image.open(out / "04_红色苹果.jpg") as apple:
                self.assertEqual(apple.size, (512, 512))
                r, g, b = apple.convert("RGB").getpixel((256, 256))
                self.assertGreater(r, 180); self.assertLess(g, 100); self.assertLess(b, 100)
                self.assertEqual(apple.convert("RGB").getpixel((0, 0)), (255, 255, 255))
                self.assertGreater(apple.convert("RGB").getpixel((330, 90))[1], 100)
            with Image.open(out / "05_蓝色方块.png") as square:
                self.assertEqual(square.size, (512, 512))
                r, g, b = square.convert("RGB").getpixel((256, 256))
                self.assertLess(r, 100); self.assertLess(g, 180); self.assertGreater(b, 180)
                self.assertEqual(square.convert("RGB").getpixel((0, 0)), (255, 255, 255))

    def test_nonempty_output_requires_force_and_force_preserves_unknown_files(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp); out.mkdir(exist_ok=True)
            (out / "unknown.txt").write_text("keep", encoding="utf-8")
            with self.assertRaises(FileExistsError): generate_demo_data(out)
            (out / "unknown.txt").unlink()
            out.rmdir()
            generate_demo_data(out)
            (out / "unknown.txt").write_text("keep", encoding="utf-8")
            generate_demo_data(out, force=True)
            self.assertEqual((out / "unknown.txt").read_text(encoding="utf-8"), "keep")

    def test_force_preserves_unknown_subdirectories(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp)
            generate_demo_data(out)
            nested = out / "unknown-subdir"
            nested.mkdir()
            marker = nested / "marker.txt"
            marker.write_text("preserve", encoding="utf-8")
            generate_demo_data(out, force=True)
            self.assertTrue(nested.is_dir())
            self.assertEqual(marker.read_text(encoding="utf-8"), "preserve")

    def test_force_rejects_foreign_or_invalid_manifest_without_modifying_files(self):
        for manifest in (
            {"schema_version": 1, "generated_by": "other.py", "files": []},
            {"schema_version": 2, "generated_by": "tools/demo/generate_demo_data.py", "files": []},
            {"schema_version": 1, "generated_by": "tools/demo/generate_demo_data.py", "files": [{"name": EXPECTED_FILES[0]}]},
            {"schema_version": 1, "generated_by": "tools/demo/generate_demo_data.py", "files": [{"name": name} for name in EXPECTED_FILES[:-1]]},
        ):
            with tempfile.TemporaryDirectory() as tmp:
                out = Path(tmp); sentinel = out / EXPECTED_FILES[0]
                sentinel.write_text("sentinel", encoding="utf-8")
                (out / "MANIFEST.json").write_text(json.dumps(manifest), encoding="utf-8")
                with self.assertRaises(FileExistsError): generate_demo_data(out, force=True)
                self.assertEqual(sentinel.read_text(encoding="utf-8"), "sentinel")

    def test_cli_generates_six_artifacts(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "cli"
            script = Path(__file__).resolve().parents[1] / "generate_demo_data.py"
            completed = subprocess.run([sys.executable, str(script), str(output)], capture_output=True, text=True)
            self.assertEqual(completed.returncode, 0, completed.stderr)
            self.assertIn(f"Generated five demo files in {output.resolve()}", completed.stdout)
            self.assertEqual(len(list(output.iterdir())), 6)

    def test_hardlink_target_is_not_modified_by_force(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; root.mkdir()
            sentinel = Path(tmp) / "sentinel.txt"; sentinel.write_text("outside", encoding="utf-8")
            generate_demo_data(root)
            target = root / EXPECTED_FILES[0]; target.unlink(); os.link(sentinel, target)
            generate_demo_data(root, force=True)
            self.assertEqual(sentinel.read_text(encoding="utf-8"), "outside")

    def test_generation_failure_is_atomic_and_retryable(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"
            with mock.patch("tools.demo.generate_demo_data._write_docx", side_effect=OSError("boom")):
                with self.assertRaises(OSError): generate_demo_data(root)
            self.assertFalse(root.exists() and list(root.iterdir()))
            generate_demo_data(root)
            old = (root / EXPECTED_FILES[0]).read_bytes(); old_manifest = (root / "MANIFEST.json").read_bytes()
            with mock.patch("tools.demo.generate_demo_data._write_docx", side_effect=OSError("boom")):
                with self.assertRaises(OSError): generate_demo_data(root, force=True)
            self.assertEqual((root / EXPECTED_FILES[0]).read_bytes(), old)
            self.assertEqual((root / "MANIFEST.json").read_bytes(), old_manifest)

    def test_manifest_is_strict_and_return_isolated(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; first = generate_demo_data(root)
            first["files"][0]["query"] = "mutated"
            second = generate_demo_data(root, force=True)
            self.assertEqual(second["files"][0]["query"], "星桥检索协议")
            bad = json.loads((root / "MANIFEST.json").read_text(encoding="utf-8")); bad["files"][0].pop("query")
            (root / "MANIFEST.json").write_text(json.dumps(bad), encoding="utf-8")
            with self.assertRaises(FileExistsError): generate_demo_data(root, force=True)

    def test_publish_failure_restores_all_old_artifacts(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; generate_demo_data(root)
            old = {p.name: p.read_bytes() for p in root.iterdir() if p.is_file()}
            original = os.replace; calls = 0
            def flaky(src, dst):
                nonlocal calls
                calls += 1
                if calls == 9: raise OSError("publish fault")
                return original(src, dst)
            with mock.patch("tools.demo.generate_demo_data.os.replace", side_effect=flaky):
                with self.assertRaises(OSError): generate_demo_data(root, force=True)
            self.assertEqual({p.name: p.read_bytes() for p in root.iterdir() if p.is_file()}, old)

    def test_keyboard_interrupt_restores_old_artifacts_and_unknowns(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; generate_demo_data(root)
            unknown = root / "unknown.txt"; unknown.write_text("keep", encoding="utf-8")
            old = {p.name: p.read_bytes() for p in root.iterdir() if p.is_file() and p.name != "unknown.txt"}
            original = os.replace; calls = 0
            def interrupting(src, dst):
                nonlocal calls
                calls += 1
                if calls == 8: raise KeyboardInterrupt()
                return original(src, dst)
            with mock.patch("tools.demo.generate_demo_data.os.replace", side_effect=interrupting):
                with self.assertRaises(KeyboardInterrupt): generate_demo_data(root, force=True)
            self.assertEqual({p.name: p.read_bytes() for p in root.iterdir() if p.is_file() and p.name != "unknown.txt"}, old)
            self.assertEqual(unknown.read_text(encoding="utf-8"), "keep")

    def test_interrupt_after_replace_before_bookkeeping_restores_everything(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; generate_demo_data(root)
            unknown = root / "unknown.txt"; unknown.write_text("keep", encoding="utf-8")
            old = {p.name: p.read_bytes() for p in root.iterdir() if p.is_file()}
            original = os.replace; calls = 0
            def replace_then_interrupt(src, dst):
                nonlocal calls
                calls += 1
                result = original(src, dst)
                if calls == 2: raise KeyboardInterrupt()
                return result
            with mock.patch("tools.demo.generate_demo_data.os.replace", side_effect=replace_then_interrupt):
                with self.assertRaises(KeyboardInterrupt): generate_demo_data(root, force=True)
            self.assertEqual({p.name: p.read_bytes() for p in root.iterdir() if p.is_file()}, old)
            self.assertEqual(unknown.read_text(encoding="utf-8"), "keep")
            self.assertTrue(json.loads((root / "MANIFEST.json").read_text(encoding="utf-8"))["schema_version"] == 1)

    def test_failed_recovery_retains_backup_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; generate_demo_data(root)
            original = os.replace; calls = 0
            def doubly_failing(src, dst):
                nonlocal calls
                calls += 1
                if calls == 8: raise KeyboardInterrupt()
                if calls == 9: raise OSError("restore fault")
                return original(src, dst)
            with mock.patch("tools.demo.generate_demo_data.os.replace", side_effect=doubly_failing):
                with self.assertRaises(RuntimeError) as ctx: generate_demo_data(root, force=True)
            self.assertIn("recovery", str(ctx.exception).lower())
            self.assertTrue(any("rollback" in p.name for p in root.parent.iterdir()))

    def test_manifest_publish_interrupt_invalidates_new_manifest_before_failed_recovery(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp) / "out"; generate_demo_data(root)
            original = os.replace; calls = 0
            def double_fault(src, dst):
                nonlocal calls
                calls += 1
                result = original(src, dst)
                if calls == 12: raise KeyboardInterrupt()
                if calls == 13: raise OSError("restore fault")
                return result
            with mock.patch("tools.demo.generate_demo_data.os.replace", side_effect=double_fault):
                with self.assertRaises(RuntimeError) as ctx: generate_demo_data(root, force=True)
            self.assertIn("recovery", str(ctx.exception).lower())
            manifest = root / "MANIFEST.json"
            self.assertFalse(manifest.exists())
            backups = [p for p in root.parent.iterdir() if "rollback" in p.name]
            self.assertTrue(backups)
            self.assertTrue((backups[0] / "MANIFEST.json").is_file())

    def test_demo_tool_declares_reproducible_dependencies(self):
        config = tomllib.loads((Path(__file__).resolve().parents[1] / "pyproject.toml").read_text(encoding="utf-8"))
        deps = " ".join(config["project"]["dependencies"]).lower()
        for name in ("python-docx", "reportlab", "pillow", "pypdfium2"):
            self.assertIn(name, deps)


if __name__ == "__main__":
    unittest.main()
